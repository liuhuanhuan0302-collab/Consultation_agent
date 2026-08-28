"""Report contact settings, snapshot compatibility, and admin authorization."""

import asyncio
import json
from io import BytesIO
from types import SimpleNamespace

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.service.reporting as reporting
from app.api.v1.endpoints import admin
from app.db.database import Base, get_db
from app.models.user import Role, User
from app.schemas.system_setting import ReportContactSettingsUpdate
from app.service.lead_export_service import generate_customer_report_docx
from app.service.pdf_service import validate_report_html
from app.service.reporting import (
    build_deepseek_prompt,
    build_report_payload,
    render_structured_report_html,
    report_data_validation_errors,
)
from app.service.system_setting_service import report_contact_snapshot, update_report_contact_settings
from app.utils.security import create_access_token


def _bearer(user_id: int) -> dict[str, str]:
    return {"Authorization": f"Bearer {create_access_token(str(user_id))}"}


def test_report_contact_settings_api_is_admin_only_and_persists_trimmed_values() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    session_factory = sessionmaker(bind=engine)
    with Session(engine) as db:
        admin_user = User(email="admin@example.com", name="Admin", role=Role.admin.value, password_hash="hash")
        sales_user = User(email="sales@example.com", name="Sales", role=Role.sales.value, password_hash="hash")
        db.add_all([admin_user, sales_user])
        db.commit()
        admin_id, sales_id = admin_user.id, sales_user.id

    app = FastAPI()
    app.include_router(admin.router)

    def override_db():
        db = session_factory()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_db
    with TestClient(app) as client:
        assert client.get("/api/admin/system-settings/report-contact").status_code == 401
        assert client.get(
            "/api/admin/system-settings/report-contact",
            headers=_bearer(sales_id),
        ).status_code == 403
        assert client.put(
            "/api/admin/system-settings/report-contact",
            headers=_bearer(sales_id),
            json={"contact_name": "越权修改"},
        ).status_code == 403

        empty = client.get(
            "/api/admin/system-settings/report-contact",
            headers=_bearer(admin_id),
        )
        assert empty.status_code == 200
        assert {key: empty.json()[key] for key in ("contact_name", "phone", "wechat", "email")} == {
            "contact_name": "",
            "phone": "",
            "wechat": "",
            "email": "",
        }

        response = client.put(
            "/api/admin/system-settings/report-contact",
            headers=_bearer(admin_id),
            json={
                "contact_name": "  优小越  ",
                "phone": " 13490000000 ",
                "wechat": "",
                "email": " contact@example.com ",
            },
        )
        assert response.status_code == 200
        assert response.json()["contact_name"] == "优小越"
        assert response.json()["wechat"] == ""
        assert response.json()["updated_by"] == "admin@example.com"

        saved = client.get(
            "/api/admin/system-settings/report-contact",
            headers=_bearer(admin_id),
        )
        assert saved.status_code == 200
        assert saved.json()["phone"] == "13490000000"
        assert saved.json()["email"] == "contact@example.com"
    engine.dispose()


def _report_data() -> dict:
    return {
        "executive_summary": [{"finding": "发现", "evidence": "Q1 得分 1/4"}],
        "dimension_analysis": [
            {
                "module_code": "M01",
                "module_name": "客户中心",
                "core_conclusion": "该维度仍需改善。",
                "evidence_rows": [{"evidence": "Q1 得分 1/4", "interpretation": "客户机制尚未闭环。"}],
            }
        ],
        "key_contradictions": [{"contradiction": "目标与机制", "evidence": "Q1", "diagnosis": "需进一步核验"}],
        "workshop_topics": [
            {"priority": "P0", "topic": "验证需求", "question": "真实需求是什么？", "deliverable": "事实清单", "nature": "形成共识"}
        ],
        "ai_scenarios": [{"name": "知识助手", "description": "检索内部资料。", "benefit": "减少查找时间。"}],
        # A model may still return a legacy extra field. V2 validation and
        # rendering must ignore it rather than exposing it to the customer.
        "management_actions": ["不应出现在新版报告"],
    }


def _v2_payload(contact: dict[str, str] | None = None) -> dict:
    return {
        "report_format_version": 2,
        "report_contact": contact or {},
        "company": {"name": "优越信息技术（广州）有限公司"},
        "score": {"total": 1, "max_score": 4, "score_rate": 0.25},
        "dimensions": [
            {"module_code": "M01", "module_name": "客户中心", "raw_score": 1, "max_score": 4, "score_rate": 0.25}
        ],
        "low_dimensions": [],
        "question_scores": [],
        "core_findings": [],
        "cases": [],
    }


def test_v2_report_uses_exact_advisory_wording_and_optional_contact_snapshot() -> None:
    payload = _v2_payload(
        {"contact_name": "优小越", "phone": "13490000000", "email": "contact@example.com"}
    )
    html = render_structured_report_html(payload, _report_data())

    assert "五、优先 AI 场景建议" in html
    assert "以下场景仅供决策参考，具体需入企调研后给出更详细的建议。" in html
    assert "六、管理层行动建议" not in html
    assert "不应出现在新版报告" not in html
    assert "进一步沟通" not in html
    assert "如需入企调研或进一步了解，可以联系：" in html
    assert "优小越" in html and "13490000000" in html and "contact@example.com" in html
    assert "微信号" not in html
    assert report_data_validation_errors({key: value for key, value in _report_data().items() if key != "management_actions"}, payload) == []

    no_contact_html = render_structured_report_html(_v2_payload(), _report_data())
    assert "进一步沟通" not in no_contact_html


def test_report_payload_snapshots_settings_and_prompt_excludes_renderer_metadata() -> None:
    lead = SimpleNamespace(
        company_name="示例公司",
        city="广州",
        industry="软件",
        company_size="1-20人",
        position="总经理",
        ai_focus="知识管理",
        lead_level="high",
        demand_summary="提升检索效率",
    )
    report = SimpleNamespace(
        submission=SimpleNamespace(total_score=1, max_score=4, score_rate=0.25, answers=[])
    )
    contact = {"contact_name": "原联系人", "phone": "10086"}
    payload = build_report_payload(lead, report, [], [], report_contact=contact)
    contact["contact_name"] = "后来修改的人"

    assert payload["report_format_version"] == 2
    assert payload["report_contact"] == {"contact_name": "原联系人", "phone": "10086"}
    prompt = build_deepseek_prompt(payload)
    assert '"management_actions"' not in prompt
    assert '"report_contact"' not in prompt
    assert "原联系人" not in prompt


def test_generation_persists_current_contact_snapshot(monkeypatch) -> None:
    lead = SimpleNamespace(
        id=3,
        company_name="示例公司",
        city="广州",
        industry="软件",
        company_size="1-20人",
        position="总经理",
        ai_focus="知识管理",
        lead_level="high",
        demand_summary="提升检索效率",
    )
    submission = SimpleNamespace(
        id=4,
        lead_id=lead.id,
        lead=lead,
        total_score=1,
        max_score=4,
        score_rate=0.25,
        answers=[],
        dimension_scores=[],
    )
    report = SimpleNamespace(
        id=5,
        submission_id=submission.id,
        submission=submission,
        company_research_json=None,
        recommendations=[],
        summary_json=None,
        html_content="",
    )

    class FakeQuery:
        def filter(self, *args, **kwargs):
            return self

        def delete(self):
            return 0

    class FakeDb:
        def __init__(self):
            self.added = []

        def query(self, *args, **kwargs):
            return FakeQuery()

        def add(self, item):
            self.added.append(item)

        def flush(self):
            return None

    async def fake_call_deepseek(payload, llm_override=None):
        return json.dumps(_report_data(), ensure_ascii=False)

    db = FakeDb()
    monkeypatch.setattr(reporting, "select_recommendations", lambda *args, **kwargs: [])
    monkeypatch.setattr(reporting, "effective_llm_override", lambda _db: None)
    monkeypatch.setattr(reporting, "report_contact_snapshot", lambda _db: {"contact_name": "生成时联系人"})
    monkeypatch.setattr(reporting, "call_deepseek", fake_call_deepseek)

    asyncio.run(reporting.generate_report_content(db, report))
    persisted = json.loads(report.summary_json)

    assert persisted["report_format_version"] == 2
    assert persisted["report_contact"] == {"contact_name": "生成时联系人"}
    assert "生成时联系人" in report.html_content
    assert "六、管理层行动建议" not in report.html_content


def test_contact_service_omits_empty_fields_and_existing_html_is_immutable() -> None:
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    with Session(engine) as db:
        update_report_contact_settings(
            db,
            ReportContactSettingsUpdate(contact_name="初始联系人", phone="", wechat="wx-1", email=""),
            updated_by="admin@example.com",
        )
        snapshot = report_contact_snapshot(db)
        persisted_payload = _v2_payload(snapshot)
        persisted_html = render_structured_report_html(persisted_payload, _report_data())

        update_report_contact_settings(
            db,
            ReportContactSettingsUpdate(contact_name="新联系人", phone="200", wechat="", email=""),
            updated_by="admin@example.com",
        )
        assert report_contact_snapshot(db) == {"contact_name": "新联系人", "phone": "200"}
        assert "初始联系人" in persisted_html
        assert "新联系人" not in persisted_html
        assert snapshot == {"contact_name": "初始联系人", "wechat": "wx-1"}

        update_report_contact_settings(
            db,
            ReportContactSettingsUpdate(),
            updated_by="admin@example.com",
        )
        assert report_contact_snapshot(db) == {}
    engine.dispose()


def test_v2_customer_word_uses_same_persisted_snapshot() -> None:
    payload = _v2_payload({"contact_name": "优小越", "phone": "13490000000", "wechat": "wx-001"})
    html = render_structured_report_html(payload, _report_data())
    report = SimpleNamespace(
        id=9,
        title="优越信息技术（广州）有限公司 AI 原生转型诊断报告",
        created_at=None,
        summary_json=json.dumps(payload, ensure_ascii=False),
        html_content=html,
    )

    validate_report_html(report)
    content = generate_customer_report_docx(report, "优越信息技术（广州）有限公司")
    document = Document(BytesIO(content))
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    assert "五、优先 AI 场景建议" in text
    assert "进一步沟通" not in text
    assert "如需入企调研或进一步了解，可以联系：" in text
    assert "优小越" in text and "13490000000" in text and "wx-001" in text
    assert "六、管理层行动建议" not in text


def test_legacy_report_rendering_keeps_historical_management_actions() -> None:
    payload = _v2_payload()
    payload.pop("report_format_version")
    html = render_structured_report_html(payload, _report_data())

    assert "五、优先 AI 场景与案例" in html
    assert "六、管理层行动建议" in html
    assert "不应出现在新版报告" in html
