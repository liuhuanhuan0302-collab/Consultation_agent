"""客户版 DOCX / Word→PDF 管道的测试：统一版式、评分头部与严格转换。"""

import asyncio
import json
import os
import subprocess
import sys
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

from app.core.config import Settings
from app.service import pdf_service
from app.service.lead_export_service import (
    BLUE,
    CUSTOMER_BODY_H1_SIZE_PT,
    CUSTOMER_BODY_LEAD_SIZE_PT,
    CUSTOMER_BODY_OVERRIDE,
    CUSTOMER_COVER_META_LABEL_WIDTH_CM,
    CUSTOMER_COVER_META_VALUE_WIDTH_CM,
    CUSTOMER_COVER_META_WIDTH_CM,
    CUSTOMER_COVER_OVERRIDE,
    DOCUMENT_STYLE_PRESET,
    NAVY,
    PALE_GRAY,
    RED,
    _chart_dimension_label,
    generate_customer_report_docx,
    generate_lead_export_docx,
)


def _customer_report() -> SimpleNamespace:
    summary = {
        "report_format_version": 2,
        "report_contact": {
            "contact_name": "优小越",
            "phone": "17646848610",
            "wechat": "18664874363",
            "email": "youxiaoyue@youkunai.cn",
        },
        "score": {"total": 106, "max_score": 242, "score_rate": 106 / 242},
        "dimensions": [
            {"module_code": "M01", "module_name": "以用户/客户为中心", "score_rate": 0.25},
            {"module_code": "M02", "module_name": "简化业务", "score_rate": 0.61},
            {"module_code": "M03", "module_name": "流程化", "score_rate": 0.58},
        ],
    }
    html = (
        '<h2>一、执行摘要</h2><p class="report-lead">优先建立可验证的场景闭环。</p>'
        "<p>摘要内容。</p>"
        '<div class="report-diagnosis-callout"><p>诊断判断：先补基础，再验证复制。</p></div>'
        '<h2>二、能力成熟度分析</h2><p class="report-section-note">以下按模块逐项分析（含该模块题目得分明细），与上方雷达图、得分排行一一对应。</p>'
        '<div class="report-module-head">M01 以用户/客户为中心 · 得分率 25%</div>'
        '<table class="report-finding-table"><thead><tr><th>序号</th><th>核心发现</th><th>证据与含义</th></tr></thead>'
        "<tbody><tr><td>1</td><td>流程需要优化</td><td>Q1 得分较低</td></tr></tbody></table>"
        "<h2>三、关键矛盾与核心诊断</h2><p>矛盾内容。</p>"
        "<h2>四、工作坊议题地图</h2><p>议题内容。</p>"
        '<h2>五、优先 AI 场景建议</h2><p class="report-section-note">旧场景说明。</p>'
        '<section class="report-case"><h4>智能派单</h4><p>场景内容。</p><p><strong>预期收益：</strong>提升效率。</p></section>'
        '<section class="report-contact-section"><h2>进一步沟通</h2><p>旧联系内容。</p></section>'
        "<h2>六、管理层行动建议</h2><ol><li>建立数据治理机制</li></ol>"
    )
    return SimpleNamespace(
        id=7,
        title="奥飞娱乐股份有限公司 AI 原生转型诊断报告",
        created_at=datetime(2026, 8, 23),
        summary_json=json.dumps(summary, ensure_ascii=False),
        html_content=html,
    )


def _docx_text(document: Document) -> str:
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    return f"{paragraphs}\n{tables}"


def _tiny_pdf() -> bytes:
    buffer = BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=A4)
    document.build([Paragraph("test", getSampleStyleSheet()["Normal"])])
    return buffer.getvalue()


def test_customer_docx_excludes_internal_lead_and_research_fields() -> None:
    docx_bytes = generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")
    document = Document(BytesIO(docx_bytes))
    text = _docx_text(document)

    assert "奥飞娱乐股份有限公司" in text
    assert "奥飞娱乐 AI 原生转型" in text
    assert "AI 原生企业转型诊断报告" in text
    assert "报告编号" in text and "RPT-000007" in text
    assert "出具日期" in text and "2026 年 8 月 23 日" in text
    for section in (
        "一、执行摘要",
        "二、能力成熟度分析",
        "三、关键矛盾与核心诊断",
        "四、工作坊议题地图",
        "五、优先 AI 场景建议",
    ):
        assert section in text
    assert "六、管理层行动建议" not in text
    assert "管理层" not in text
    assert "进一步沟通" not in text
    assert "如需入企调研或进一步了解，可以联系" in text
    for internal in (
        "一、基本信息",
        "二、企业公开信息与 AI 情报分析",
        "手机号",
        "首次查看",
        "AI 情报",
        "公司介绍",
        "信息来源",
        "后台",
    ):
        assert internal not in text
    assert len(document.inline_shapes) == 2


def test_customer_docx_score_header_uses_summary_fields() -> None:
    docx_bytes = generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")
    text = _docx_text(Document(BytesIO(docx_bytes)))

    assert "诊断得分" in text and "106 / 242" in text
    assert "综合得分率" in text and "44%" in text

    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    cover_footer_at = document_xml.index("让 AI 从局部工具走向企业级生产力")
    page_break_at = document_xml.index('w:type="page"', cover_footer_at)
    score_at = document_xml.index("诊断结果概览")
    assert cover_footer_at < page_break_at < score_at


def test_m01_chart_label_is_concise_while_report_name_stays_unchanged() -> None:
    dimension = {
        "module_code": "M01",
        "module_name": "一心：以用户/客户为中心",
    }

    assert _chart_dimension_label(dimension) == "用户/客户中心"
    assert dimension["module_name"] == "一心：以用户/客户为中心"


def test_customer_docx_uses_executive_consulting_visual_contract() -> None:
    docx_bytes = generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")
    document = Document(BytesIO(docx_bytes))
    text = _docx_text(document)

    assert "从诊断共识走向可执行的 AI 转型路径" in text
    assert "奥飞娱乐 AI 原生转型" in text and "诊断报告" in text
    assert "让 AI 从局部工具走向企业级生产力" in text
    assert DOCUMENT_STYLE_PRESET == "standard_business_brief"
    assert CUSTOMER_COVER_OVERRIDE == "approved_reference_editorial_cover"
    assert CUSTOMER_BODY_OVERRIDE == "reference_consulting_body_v2"
    assert abs(document.sections[0].page_width - Cm(21)) < 1000
    assert abs(document.sections[0].page_height - Cm(29.7)) < 1000
    assert document.sections[0].different_first_page_header_footer is True
    first_page_header = document.sections[0].first_page_header.paragraphs[0]
    first_page_properties = first_page_header._p.get_or_add_pPr()
    assert first_page_properties.find(qn("w:pBdr")) is None
    footer_text = "\n".join(
        cell.text
        for table in document.sections[0].footer.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "企业 AI 转型诊断" in footer_text
    assert "2026 年 8 月 23 日" in footer_text
    header = document.sections[0].header.paragraphs[0]
    paragraph_borders = header._p.pPr.find(qn("w:pBdr"))
    assert paragraph_borders is not None
    bottom_border = paragraph_borders.find(qn("w:bottom"))
    assert bottom_border is not None
    assert bottom_border.get(qn("w:color")) == RED

    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert f'w:fill="{NAVY}"' in document_xml
    assert f'w:fill="{PALE_GRAY}"' in document_xml
    assert f'w:color="{RED}"' in document_xml
    for forbidden in ("保密级别", "请妥善保管", "EXECUTIVE DIAGNOSTIC"):
        assert forbidden not in text


def test_customer_docx_cover_has_five_borderless_fixed_metadata_rows() -> None:
    document = Document(BytesIO(generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")))
    cover_meta = document.tables[0]

    assert len(cover_meta.rows) == 5
    assert [row.cells[0].text for row in cover_meta.rows] == [
        "评估对象",
        "报告类型",
        "评估范围",
        "报告编号",
        "出具日期",
    ]
    assert [row.cells[1].text for row in cover_meta.rows] == [
        "奥飞娱乐股份有限公司",
        "AI 原生企业转型诊断报告",
        "企业 AI 原生能力成熟度与转型路径",
        "RPT-000007",
        "2026 年 8 月 23 日",
    ]
    table_properties = cover_meta._tbl.tblPr
    assert table_properties.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
    assert int(table_properties.find(qn("w:tblW")).get(qn("w:w"))) == Cm(CUSTOMER_COVER_META_WIDTH_CM).twips
    assert int(table_properties.find(qn("w:tblInd")).get(qn("w:w"))) == 0
    grid_widths = [int(column.get(qn("w:w"))) for column in cover_meta._tbl.tblGrid]
    assert grid_widths == [
        Cm(CUSTOMER_COVER_META_LABEL_WIDTH_CM).twips,
        Cm(CUSTOMER_COVER_META_VALUE_WIDTH_CM).twips,
    ]
    for row in cover_meta.rows:
        assert [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells] == grid_widths
        for cell in row.cells:
            margins = cell._tc.tcPr.find(qn("w:tcMar"))
            assert margins is not None
            assert int(margins.find(qn("w:top")).get(qn("w:w"))) == 55
            assert int(margins.find(qn("w:bottom")).get(qn("w:w"))) == 55
    borders = table_properties.find(qn("w:tblBorders"))
    assert borders is not None
    assert all(border.get(qn("w:val")) == "nil" for border in borders)


def test_customer_docx_reference_body_tokens_and_native_callout() -> None:
    document = Document(
        BytesIO(generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司"))
    )

    h1 = next(paragraph for paragraph in document.paragraphs if paragraph.text == "一、执行摘要")
    assert h1.runs[0].font.size.pt == CUSTOMER_BODY_H1_SIZE_PT
    assert str(h1.runs[0].font.color.rgb) == NAVY
    assert h1._p.get_or_add_pPr().find(qn("w:pBdr")) is None

    lead = next(paragraph for paragraph in document.paragraphs if paragraph.text == "优先建立可验证的场景闭环。")
    assert lead.runs[0].font.size.pt == CUSTOMER_BODY_LEAD_SIZE_PT
    assert str(lead.runs[0].font.color.rgb) == RED

    callout = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("诊断判断："))
    properties = callout._p.get_or_add_pPr()
    assert properties.find(qn("w:shd")).get(qn("w:fill")) == "FBE9E9"
    assert properties.find(qn("w:pBdr")).find(qn("w:left")).get(qn("w:color")) == RED
    assert str(callout.runs[0].font.color.rgb) == NAVY

    assert BLUE == "2F5597"


def test_customer_docx_long_short_name_uses_bounded_title_reduction() -> None:
    legal_name = "华南国际数字科技创新发展集团股份有限公司"
    document = Document(BytesIO(generate_customer_report_docx(_customer_report(), legal_name)))
    short_title = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.endswith("AI 原生转型")
    )

    assert short_title.text == "华南国际数字科技创新发展 AI 原生转型"
    assert 20.5 <= short_title.runs[0].font.size.pt <= 23.0


def test_chromium_fallback_html_is_a4_consulting_report_without_dashboard_hero() -> None:
    html = pdf_service.render_report_html_attachment(_customer_report()).decode("utf-8")

    assert "@page { size: A4" in html
    assert "page-break-after: always" in html
    assert ".report-html { background: #fff; padding:" in html
    assert "#17365D" in html and "#C00000" in html and "#2F5597" in html and "#F1F4F8" in html
    assert "report-cover" in html and "让 AI 从局部工具走向企业级生产力" in html
    assert "width: 162mm" in html and "width: 134mm" in html
    assert 'data-body-style="reference_consulting_body_v2"' in html
    assert "report-running-header" in html and "report-running-footer" in html
    cover_html = html.split('<section class="report-cover">', 1)[1].split("</section>", 1)[0]
    assert cover_html.count("<dt>") == 5
    for label in ("评估对象", "报告类型", "评估范围", "报告编号", "出具日期"):
        assert f"<dt>{label}</dt>" in cover_html
    for forbidden in ("保密级别", "请妥善保管", "EXECUTIVE DIAGNOSTIC", "cover-kicker"):
        assert forbidden not in cover_html
    assert "奥飞娱乐股份有限公司" in cover_html
    assert "奥飞娱乐 AI 原生转型" in cover_html
    assert "2026 年 8 月 23 日" in cover_html
    assert "score-strip" not in cover_html
    assert html.index('<section class="report-overview">') < html.index('class="score-strip"')
    assert html.index('class="score-strip"') < html.index('aria-label="执行摘要"')
    assert html.index('aria-label="执行摘要"') < html.index('class="chart-grid"')
    assert html.count("一、执行摘要") == 1
    assert html.index('class="chart-grid"') < html.index("二、能力成熟度分析")
    assert "linear-gradient" not in html
    assert "radial-gradient" not in html
    assert "border-radius: 18px" not in html
    assert "一、执行摘要" in html and "五、优先 AI 场景建议" in html
    assert "六、管理层行动建议" not in html and "管理层" not in html
    assert "以下场景仅供决策参考，具体需入企调研后给出更详细的建议。" in html


def test_chromium_fallback_html_excludes_internal_customer_and_research_fields() -> None:
    report = _customer_report()
    report.company_name = "奥飞娱乐"
    report.company_research_json = "RESEARCH_PRIVATE_SENTINEL"
    report.generation_error = "AUDIT_PRIVATE_SENTINEL"
    report.submission = SimpleNamespace(
        lead=SimpleNamespace(
            company_name="奥飞娱乐",
            contact_name="CONTACT_PRIVATE_SENTINEL",
            phone="PHONE_PRIVATE_SENTINEL",
            email="EMAIL_PRIVATE_SENTINEL",
            wechat="WECHAT_PRIVATE_SENTINEL",
        )
    )

    html = pdf_service.render_report_html_attachment(report).decode("utf-8")
    for sentinel in (
        "RESEARCH_PRIVATE_SENTINEL",
        "AUDIT_PRIVATE_SENTINEL",
        "CONTACT_PRIVATE_SENTINEL",
        "PHONE_PRIVATE_SENTINEL",
        "EMAIL_PRIVATE_SENTINEL",
        "WECHAT_PRIVATE_SENTINEL",
    ):
        assert sentinel not in html


@pytest.mark.parametrize(
    ("summary_json", "error_match"),
    [
        (None, "缺少 score 评分快照"),
        ("not-json", "不是有效 JSON"),
        (json.dumps({"score": {"total": True, "max_score": 100, "score_rate": 0.5}}), "非布尔数值"),
        (json.dumps({"score": {"total": float("nan"), "max_score": 100, "score_rate": 0.5}}), "有限数值"),
        (json.dumps({"score": {"total": -1, "max_score": 100, "score_rate": -0.01}}), "不得小于 0"),
        (json.dumps({"score": {"total": 101, "max_score": 100, "score_rate": 1.01}}), "不得大于"),
        (json.dumps({"score": {"total": 50, "max_score": 100, "score_rate": 0.7}}), "不一致"),
    ],
)
def test_invalid_persisted_score_blocks_all_pdf_renderers(
    monkeypatch,
    summary_json,
    error_match,
) -> None:
    report = _customer_report()
    report.summary_json = summary_json
    renderer_calls = {"docx": 0, "browser": 0}

    def forbidden_docx(_docx_bytes):
        renderer_calls["docx"] += 1
        raise AssertionError("无效评分不得进入 LibreOffice 渲染")

    def forbidden_browser(_html_bytes):
        renderer_calls["browser"] += 1
        raise AssertionError("无效评分不得进入 Chromium fallback")

    monkeypatch.setattr(pdf_service, "convert_customer_docx_to_pdf", forbidden_docx)
    monkeypatch.setattr(pdf_service, "render_report_pdf_bytes_with_browser_html", forbidden_browser)

    with pytest.raises(pdf_service.ReportPdfValidationError, match=error_match):
        asyncio.run(pdf_service.render_report_pdf_bytes(report))

    assert renderer_calls == {"docx": 0, "browser": 0}


def test_customer_docx_package_never_contains_internal_data_sentinels() -> None:
    report = _customer_report()
    report.company_research_json = json.dumps(
        {"company_overview": "RESEARCH_PRIVATE_SENTINEL", "sources": ["SEARCH_SOURCE_SENTINEL"]}
    )
    report.submission = SimpleNamespace(
        lead=SimpleNamespace(
            company_name="奥飞娱乐",
            contact_name="CONTACT_PRIVATE_SENTINEL",
            phone="PHONE_PRIVATE_SENTINEL",
            email="EMAIL_PRIVATE_SENTINEL",
            wechat="WECHAT_PRIVATE_SENTINEL",
            source_code="SOURCE_PRIVATE_SENTINEL",
            first_viewed_by="ADMIN_PRIVATE_SENTINEL",
        )
    )

    package = generate_customer_report_docx(report, "奥飞娱乐")
    with zipfile.ZipFile(BytesIO(package)) as archive:
        xml_payload = b"\n".join(
            archive.read(name)
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        ).decode("utf-8", errors="ignore")

    for sentinel in (
        "RESEARCH_PRIVATE_SENTINEL",
        "SEARCH_SOURCE_SENTINEL",
        "CONTACT_PRIVATE_SENTINEL",
        "PHONE_PRIVATE_SENTINEL",
        "EMAIL_PRIVATE_SENTINEL",
        "WECHAT_PRIVATE_SENTINEL",
        "SOURCE_PRIVATE_SENTINEL",
        "ADMIN_PRIVATE_SENTINEL",
    ):
        assert sentinel not in xml_payload


def test_customer_cover_and_score_tables_have_fixed_grid_widths() -> None:
    document = Document(BytesIO(generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")))

    assert len(document.tables) >= 3
    for table, expected_widths in zip(
        document.tables[:2],
        (
            [Cm(CUSTOMER_COVER_META_LABEL_WIDTH_CM).twips, Cm(CUSTOMER_COVER_META_VALUE_WIDTH_CM).twips],
            [Cm(17.4 / 2).twips] * 2,
        ),
    ):
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        assert layout is not None and layout.get(qn("w:type")) == "fixed"
        grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        assert grid_widths == expected_widths
        cell_widths = [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in table.rows[0].cells]
        assert cell_widths == grid_widths


def test_render_prefers_docx_when_enabled(monkeypatch) -> None:
    report = _customer_report()
    valid_pdf = _tiny_pdf()
    called = {"docx": False, "browser": False}

    def fake_docx(docx_bytes):
        assert isinstance(docx_bytes, bytes) and docx_bytes.startswith(b"PK")
        called["docx"] = True
        return valid_pdf

    def fake_browser(html_bytes):
        assert isinstance(html_bytes, bytes)
        called["browser"] = True
        return valid_pdf

    monkeypatch.setattr(pdf_service, "convert_customer_docx_to_pdf", fake_docx)
    monkeypatch.setattr(pdf_service, "render_report_pdf_bytes_with_browser_html", fake_browser)

    assert asyncio.run(pdf_service.render_report_pdf_bytes(report)) == valid_pdf
    assert called == {"docx": True, "browser": False}


def test_docx_failure_retries_three_times_and_never_calls_browser(monkeypatch) -> None:
    report = _customer_report()
    calls = {"docx": 0, "browser": 0}

    def fail_docx(_bytes):
        calls["docx"] += 1
        raise RuntimeError("LibreOffice 转换失败（模拟）")

    monkeypatch.setattr(
        pdf_service,
        "convert_customer_docx_to_pdf",
        fail_docx,
    )
    monkeypatch.setattr(
        pdf_service,
        "render_report_pdf_bytes_with_browser_html",
        lambda _html: calls.__setitem__("browser", calls["browser"] + 1),
    )

    with pytest.raises(pdf_service.CustomerPdfConversionError, match="连续失败 3 次"):
        asyncio.run(pdf_service.render_report_pdf_bytes(report))
    assert calls == {"docx": 3, "browser": 0}


def test_docx_failure_is_strict_even_if_legacy_fallback_flag_is_enabled(monkeypatch) -> None:
    report = _customer_report()
    monkeypatch.setattr(
        pdf_service,
        "get_settings",
        lambda: Settings(pdf_docx_fallback_to_browser=True, pdf_browser_render=True),
    )
    monkeypatch.setattr(
        pdf_service,
        "convert_customer_docx_to_pdf",
        lambda _bytes: (_ for _ in ()).throw(RuntimeError("LibreOffice 转换失败（模拟）")),
    )

    with pytest.raises(pdf_service.CustomerPdfConversionError, match="已转人工处理"):
        asyncio.run(pdf_service.render_report_pdf_bytes(report))


def test_docx_disabled_blocks_customer_attachment_instead_of_using_browser(monkeypatch) -> None:
    report = _customer_report()
    monkeypatch.setattr(pdf_service, "get_settings", lambda: Settings(pdf_docx_render=False))
    monkeypatch.setattr(
        pdf_service,
        "convert_customer_docx_to_pdf",
        lambda _bytes: (_ for _ in ()).throw(AssertionError("docx 渲染不应被调用")),
    )
    monkeypatch.setattr(
        pdf_service,
        "render_report_pdf_bytes_with_browser_html",
        lambda _html: (_ for _ in ()).throw(AssertionError("浏览器 fallback 不应被调用")),
    )

    with pytest.raises(pdf_service.CustomerPdfConversionError, match="仅允许 Word→PDF"):
        asyncio.run(pdf_service.render_report_pdf_bytes(report))


def test_worker_threads_receive_bytes_instead_of_orm_report(monkeypatch) -> None:
    report = _customer_report()
    valid_pdf = _tiny_pdf()
    threaded_arguments: list[tuple[object, ...]] = []

    async def fake_to_thread(function, *args, **kwargs):
        threaded_arguments.append(args)
        assert args and all(isinstance(value, bytes) for value in args)
        return function(*args, **kwargs)

    monkeypatch.setattr(pdf_service.asyncio, "to_thread", fake_to_thread)
    monkeypatch.setattr(pdf_service, "convert_customer_docx_to_pdf", lambda _docx: valid_pdf)

    assert asyncio.run(pdf_service.render_report_pdf_bytes(report)) == valid_pdf
    assert len(threaded_arguments) == 1


def test_libreoffice_failure_keeps_root_cause_on_conversion_exception(monkeypatch) -> None:
    report = _customer_report()
    monkeypatch.setattr(
        pdf_service,
        "get_settings",
        lambda: Settings(pdf_docx_fallback_to_browser=True, pdf_browser_render=False),
    )
    monkeypatch.setattr(
        pdf_service,
        "convert_customer_docx_to_pdf",
        lambda _bytes: (_ for _ in ()).throw(RuntimeError("LO_ROOT_CAUSE_SENTINEL")),
    )

    with pytest.raises(pdf_service.CustomerPdfConversionError) as captured:
        asyncio.run(pdf_service.render_report_pdf_bytes(report))
    assert captured.value.__cause__ is not None
    assert "LO_ROOT_CAUSE_SENTINEL" in str(captured.value.__cause__)


def test_libreoffice_conversion_invokes_soffice_and_returns_pdf(monkeypatch, tmp_path) -> None:
    fake_soffice = tmp_path / "soffice"
    fake_soffice.write_text("")
    monkeypatch.setattr(pdf_service, "libreoffice_executable", lambda: str(fake_soffice))
    captured: dict = {}

    def fake_run(command, **kwargs):
        captured["command"] = command
        outdir = Path(command[command.index("--outdir") + 1])
        captured["temp_root"] = outdir.parent
        assert captured["temp_root"].exists()
        assert (captured["temp_root"] / "lo-profile").is_dir()
        pdf_bytes = _tiny_pdf()
        (outdir / "customer-report.pdf").write_bytes(pdf_bytes)
        captured["pdf"] = pdf_bytes
        return SimpleNamespace(returncode=0, stdout="convert customer-report.docx as writer_pdf_Export", stderr="")

    monkeypatch.setattr(pdf_service.subprocess, "run", fake_run)

    docx_bytes = generate_customer_report_docx(_customer_report(), "奥飞娱乐股份有限公司")
    result = pdf_service.convert_customer_docx_to_pdf(docx_bytes)
    assert result == captured["pdf"]

    command = captured["command"]
    assert command[0] == str(fake_soffice)
    assert "--headless" in command
    assert "--convert-to" in command and "pdf:writer_pdf_Export" in command
    env_flag = next(arg for arg in command if arg.startswith("-env:UserInstallation="))
    assert env_flag.endswith("lo-profile")
    docx_arg = next(arg for arg in command if arg.endswith(".docx"))
    assert Path(docx_arg).name == "customer-report.docx"
    assert "--outdir" in command
    assert Path(command[command.index("--outdir") + 1]).name == "output"
    assert Path(docx_arg).parent.name == "input"
    assert not captured["temp_root"].exists()


def test_libreoffice_failure_reports_clear_error(monkeypatch, tmp_path) -> None:
    fake_soffice = tmp_path / "soffice"
    fake_soffice.write_text("")
    monkeypatch.setattr(pdf_service, "libreoffice_executable", lambda: str(fake_soffice))
    monkeypatch.setattr(
        pdf_service.subprocess,
        "run",
        lambda command, **kwargs: SimpleNamespace(
            returncode=1,
            stdout="writer stdout diagnostic",
            stderr="soffice crashed",
        ),
    )

    with pytest.raises(RuntimeError, match="stdout='writer stdout diagnostic'.*stderr='soffice crashed'"):
        pdf_service.convert_customer_docx_to_pdf(b"PK fake docx")


def test_libreoffice_timeout_reports_clear_error(monkeypatch, tmp_path) -> None:
    fake_soffice = tmp_path / "soffice"
    fake_soffice.write_text("")
    monkeypatch.setattr(pdf_service, "libreoffice_executable", lambda: str(fake_soffice))

    def fake_run(command, **kwargs):
        raise subprocess.TimeoutExpired(command, kwargs["timeout"])

    monkeypatch.setattr(pdf_service.subprocess, "run", fake_run)

    with pytest.raises(RuntimeError, match="转换超时"):
        pdf_service.convert_customer_docx_to_pdf(b"PK fake docx")


def test_libreoffice_timeout_configuration_is_bounded() -> None:
    with pytest.raises(ValueError):
        Settings(libreoffice_timeout=9)
    with pytest.raises(ValueError):
        Settings(libreoffice_timeout=601)


def test_configured_libreoffice_command_is_resolved_from_path(monkeypatch, tmp_path) -> None:
    executable = tmp_path / "custom-soffice"
    executable.write_text("")
    monkeypatch.setattr(
        pdf_service,
        "get_settings",
        lambda: Settings(libreoffice_executable="custom-soffice"),
    )
    monkeypatch.setattr(
        pdf_service.shutil,
        "which",
        lambda command: str(executable) if command == "custom-soffice" else None,
    )

    assert pdf_service.libreoffice_executable() == str(executable)


def test_invalid_configured_libreoffice_path_is_actionable(monkeypatch) -> None:
    monkeypatch.setattr(
        pdf_service,
        "get_settings",
        lambda: Settings(libreoffice_executable="missing-custom-soffice"),
    )
    monkeypatch.setattr(pdf_service.shutil, "which", lambda _command: None)

    with pytest.raises(RuntimeError, match="LIBREOFFICE_EXECUTABLE.*不可用"):
        pdf_service.libreoffice_executable()


def test_fixture_script_generates_docx_without_database_query(tmp_path) -> None:
    environment = os.environ.copy()
    environment.update(
        {
            "ENVIRONMENT": "development",
            "DATABASE_URL": "sqlite:///:memory:",
            "LIBREOFFICE_EXECUTABLE": str(tmp_path / "intentionally-missing-soffice"),
        }
    )
    script = Path(__file__).resolve().parents[1] / "scripts" / "generate_customer_report.py"
    result = subprocess.run(
        [sys.executable, "-B", str(script), "--fixture", "--outdir", str(tmp_path)],
        cwd=script.parents[1],
        env=environment,
        check=False,
        capture_output=True,
        text=True,
        timeout=60,
    )

    assert result.returncode == 0, result.stderr
    assert (tmp_path / "示例科技集团有限公司_AI诊断报告.docx").exists()
    assert not list(tmp_path.glob("*-browser-preview.*"))
    assert "已生成客户版 DOCX" in result.stdout
    assert "Chromium 预览" not in result.stdout
    assert "读取数据库失败" not in result.stdout
