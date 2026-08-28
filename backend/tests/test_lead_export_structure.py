import json
from io import BytesIO
from types import SimpleNamespace

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.service.lead_export_service import (
    FONT_NAME,
    generate_customer_report_docx,
    generate_lead_export_docx,
)


def _lead() -> SimpleNamespace:
    return SimpleNamespace(
        company_name="示例科技有限公司",
        city="青岛市",
        industry="制造业",
        company_size="100-500人",
        annual_revenue="1-5亿",
        contact_name="张经理",
        position="总经理",
        phone="13800000000",
        email="customer@example.com",
        wechat="wechat-id",
        source_code="default",
        created_at=None,
        ai_focus="提升销售与交付效率",
        demand_summary=None,
        lead_level="high",
    )


def _submission() -> SimpleNamespace:
    return SimpleNamespace(total_score=120, max_score=240, score_rate=0.5, submitted_at=None)


def _report() -> SimpleNamespace:
    research = {
        "company_overview": [
            {"title": "公司沿革与定位", "content": "公司公开介绍。[来源1]"},
            {"title": "主营业务", "content": "聚焦工业自动化。"},
        ],
        "revenue_scale": "",
        "products": "工业自动化产品。[来源1]",
        "industry_characteristics": "行业竞争激烈。[来源1]",
        "development_status": "处于业务扩张阶段。[来源1]",
        "challenges": "可能面临交付协同挑战。",
        "ai_opportunities": "可建设销售与交付智能体。",
        "analysis": "建议先完成数据治理，再建设关键场景。",
        "sources": [{"title": "企业官方网站", "url": "https://example.com"}],
        "source_refs": {"company_overview": [1], "products": [1]},
        "researched_at": "2026-08-22T14:00:00",
    }
    summary = {
        "report_format_version": 2,
        "report_contact": {
            "contact_name": "优小越",
            "phone": "17646848610",
            "wechat": "18664874363",
            "email": "youxiaoyue@youkunai.cn",
        },
        "score": {"total": 120, "max_score": 240, "score_rate": 0.5},
        "dimensions": [
            {"module_code": "M01", "module_name": "以用户/客户为中心", "score_rate": 0.25},
            {"module_code": "M02", "module_name": "简化业务", "score_rate": 0.61},
            {"module_code": "M03", "module_name": "流程化", "score_rate": 0.58},
            {"module_code": "M04", "module_name": "数字化", "score_rate": 0.32},
            {"module_code": "M05", "module_name": "智能化", "score_rate": 0.32},
        ]
    }
    html = """
    <article class="report-document">
      <section>
        <h2>一、执行摘要</h2>
        <table class="report-finding-table"><thead><tr><th>序号</th><th>核心发现</th><th>证据与含义</th></tr></thead>
        <tbody><tr><td>1</td><td><strong>流程需要优化</strong></td><td>Q1 得分较低</td></tr></tbody></table>
      </section>
      <section>
        <h2>二、能力成熟度分析</h2><p class="report-section-note">以下按模块逐项分析（含该模块题目得分明细），与上方雷达图、得分排行一一对应。</p>
        <div class="report-module-head">M01 以用户/客户为中心 · 得分率 25%</div>
        <table class="report-finding-table report-cad-table">
          <thead><tr><th>核心结论</th><th>数据依据</th><th>分析解读</th></tr></thead>
          <tbody>
            <tr><td>该维度需要优先改进。</td><td>Q1 得分较低</td><td>客户机制不完整。</td></tr>
            <tr><td></td><td>Q2 得分一般</td><td>数据闭环尚未建立。</td></tr>
          </tbody>
        </table>
      </section>
      <section><h2>三、关键矛盾与核心诊断</h2>
        <table class="report-finding-table report-contradiction-table"><thead><tr><th>关键矛盾</th><th>证据</th><th>诊断</th></tr></thead>
        <tbody><tr><td>局部清晰与整体协同</td><td>M01 与 M04 得分差异明显。</td><td>需要统一跨部门机制。</td></tr></tbody></table>
      </section>
      <section><h2>四、工作坊议题地图</h2><p>围绕关键机制形成共识。</p></section>
      <section><h2>五、优先 AI 场景建议</h2><p class="report-section-note">旧场景说明。</p>
        <section class="report-case"><h4>智能排产</h4><p>结合订单预测优化排产。</p><p><strong>预期收益：</strong>提升交付效率。</p></section>
      </section>
      <section class="report-contact-section"><h2>进一步沟通</h2><p>旧联系信息。</p></section>
      <section><h2>六、管理层行动建议</h2><ol><li>建立数据治理机制</li></ol></section>
    </article>
    """
    return SimpleNamespace(
        company_research_json=json.dumps(research, ensure_ascii=False),
        summary_json=json.dumps(summary, ensure_ascii=False),
        html_content=html,
    )


def test_lead_export_has_three_paginated_parts_and_native_report_tables() -> None:
    content = generate_lead_export_docx(_lead(), _submission(), _report(), source_name="官网", final_report_sent=True)
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "一、基本信息" in text
    assert "二、企业公开信息与 AI 情报分析" in text
    assert "三、客户最终诊断报告" in text
    assert "1. 公司介绍" in text
    assert "公司沿革与定位" in text
    assert "主营业务" in text
    assert "资料来源：[来源1]" in text
    assert "见本部分信息来源清单" not in text
    assert "资料来源：证据校验失败，未找到可核验来源编号" in text
    assert "2. 营收规模" in text
    assert "暂未检索到可靠公开信息" in text
    assert "8. AI 综合分析" in text
    assert "一、执行摘要" in text
    assert "三、关键矛盾与核心诊断" in text
    assert "四、工作坊议题地图" in text
    assert "五、优先 AI 场景建议" in text
    assert "六、管理层行动建议" not in text
    assert "管理层" not in text
    assert "如需入企调研或进一步了解，可以联系" in text
    assert "进一步沟通" not in text
    summary_table = next(table for table in document.tables if table.cell(0, 0).text == "序号")
    assert summary_table.cell(0, 1).text == "核心发现"
    first_width = int(summary_table.cell(0, 0)._tc.tcPr.tcW.get(qn("w:w")))
    evidence_width = int(summary_table.cell(0, 2)._tc.tcPr.tcW.get(qn("w:w")))
    summary_grid_widths = [int(column.get(qn("w:w"))) for column in summary_table._tbl.tblGrid]
    assert 650 <= first_width < 750
    assert evidence_width > first_width * 8
    assert summary_grid_widths[0] == first_width
    assert summary_grid_widths[2] == evidence_width
    assert len(document.inline_shapes) == 2
    report_table = next(table for table in document.tables if table.cell(0, 0).text == "核心结论")
    assert report_table._tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
    assert report_table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
    assert report_table.cell(1, 0)._tc is report_table.cell(2, 0)._tc
    assert report_table.cell(1, 0).text == "该维度需要优先改进。"
    assert report_table.cell(1, 0).vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
    assert report_table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    report_grid_widths = [int(column.get(qn("w:w"))) for column in report_table._tbl.tblGrid]
    assert report_grid_widths[0] < report_grid_widths[1] / 3
    assert report_grid_widths[2] > report_grid_widths[1]
    assert len(document.element.xpath('.//w:br[@w:type="page"]')) >= 3
    assert document.styles["Normal"].font.name == FONT_NAME


def test_lead_export_does_not_include_unsent_report_snapshot() -> None:
    content = generate_lead_export_docx(_lead(), _submission(), _report(), final_report_sent=False)
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "客户最终诊断报告尚未发送" in text
    assert "一、执行摘要" not in text


def _final_report_layout_signature(document: Document) -> tuple:
    report_tables = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text in {"序号", "核心结论"}
    ]
    table_signatures = []
    for table in report_tables:
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        header_shading = table.rows[0].cells[0]._tc.tcPr.find(qn("w:shd"))
        table_signatures.append(
            (
                tuple(int(column.get(qn("w:w"))) for column in table._tbl.tblGrid),
                layout.get(qn("w:type")) if layout is not None else None,
                header_properties.find(qn("w:tblHeader")) is not None,
                header_shading.get(qn("w:fill")) if header_shading is not None else None,
                tuple(
                    (
                        run.font.name,
                        run.font.size.pt if run.font.size else None,
                        run.bold,
                    )
                    for run in table.rows[0].cells[0].paragraphs[0].runs
                ),
            )
        )
    heading_signatures = []
    for paragraph in document.paragraphs:
        if paragraph.text in {"一、执行摘要", "二、能力成熟度分析", "五、优先 AI 场景建议"}:
            run = paragraph.runs[0]
            heading_signatures.append(
                (
                    paragraph.text,
                    run.font.name,
                    run.font.size.pt if run.font.size else None,
                    str(run.font.color.rgb),
                    run.bold,
                    paragraph.paragraph_format.keep_with_next,
                )
            )
    return tuple(table_signatures), tuple(heading_signatures), len(document.inline_shapes)


def test_internal_part_three_and_customer_docx_share_final_layout_signature() -> None:
    report = _report()
    internal = Document(
        BytesIO(generate_lead_export_docx(_lead(), _submission(), report, final_report_sent=True))
    )
    customer = Document(BytesIO(generate_customer_report_docx(report, "示例科技有限公司")))

    assert _final_report_layout_signature(internal) == _final_report_layout_signature(customer)
