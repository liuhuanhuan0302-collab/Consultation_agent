"""生成后台单客户 Word 档案。"""

from __future__ import annotations

import re
from html import unescape
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models import CompanyLead, DiagnosisSubmission, Report
from app.utils.time_utils import to_china_time, utc_now


FONT_NAME = "Microsoft YaHei"


def _set_run_font(run) -> None:
    """同时写入西文字体与东亚字体，避免中文在 Word 中回退到默认宋体。"""
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_paragraph_font(paragraph) -> None:
    for run in paragraph.runs:
        _set_run_font(run)


def _set_document_font(document: Document) -> None:
    for style_name in ("Normal", "Title", "Subtitle"):
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _text(value: object | None, fallback: str = "未填写") -> str:
    if value is None or str(value).strip() == "":
        return fallback
    return str(value).strip()


def _format_datetime(value: datetime | None) -> str:
    if value is None:
        return "未提交"
    return to_china_time(value).strftime("%Y-%m-%d %H:%M:%S")


def _report_text(html: str | None) -> str:
    if not html:
        return "报告暂未生成。"
    value = re.sub(r"<\s*br\s*/?\s*>", "\n", html, flags=re.IGNORECASE)
    value = re.sub(r"</(?:p|div|section|h[1-6]|li|tr)\s*>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"<[^>]+>", "", value)
    value = unescape(value).replace("\u00a0", " ")
    return re.sub(r"\n{3,}", "\n\n", value).strip()


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _set_cell_text(cell, label: str, value: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _shade(cell, "F5F8FC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}\n")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(92, 108, 132)
    _set_run_font(label_run)
    value_run = paragraph.add_run(value)
    value_run.bold = True
    value_run.font.size = Pt(10.5)
    _set_run_font(value_run)


def _set_plain_cell_text(cell, value: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.text = value
    paragraph.paragraph_format.space_after = Pt(0)
    _set_paragraph_font(paragraph)


def _add_section_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(20, 76, 112)
    _set_run_font(run)


def generate_lead_export_docx(
    lead: CompanyLead,
    submission: DiagnosisSubmission | None,
    report: Report | None,
    source_name: str | None = None,
) -> bytes:
    """返回一个包含客户资料、诊断数据和报告内容的 .docx 文件。"""
    document = Document()
    _set_document_font(document)
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{_text(lead.company_name, '客户')} AI 转型诊断档案")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(15, 56, 93)
    _set_run_font(title_run)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"导出时间：{_format_datetime(utc_now())}（北京时间）")
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = RGBColor(102, 112, 133)
    _set_run_font(subtitle_run)

    _add_section_title(document, "一、客户基本信息")
    profile = [
        ("公司", _text(lead.company_name)),
        ("行业", _text(lead.industry)),
        ("企业规模", _text(lead.company_size)),
        ("年营收", _text(lead.annual_revenue)),
        ("联系人", _text(lead.contact_name)),
        ("职位", _text(lead.position)),
        ("手机号", _text(lead.phone)),
        ("邮箱", _text(lead.email)),
        ("微信", _text(lead.wechat)),
        ("来源", _text(source_name or lead.source_code)),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index in range(0, len(profile), 2):
        row = table.add_row().cells
        _set_cell_text(row[0], *profile[index])
        _set_cell_text(row[1], *profile[index + 1])

    focus = document.add_paragraph()
    focus.paragraph_format.space_before = Pt(8)
    focus_label = focus.add_run("AI 转型关注点：")
    focus_label.bold = True
    _set_run_font(focus_label)
    focus_value = focus.add_run(_text(lead.ai_focus or lead.demand_summary, "暂未填写"))
    _set_run_font(focus_value)

    _add_section_title(document, "二、诊断结果")
    if not submission:
        paragraph = document.add_paragraph("客户尚未完成诊断问卷。")
        _set_paragraph_font(paragraph)
    else:
        result = [
            ("线索等级", _text(lead.lead_level, "LOW").upper()),
            ("综合得分", f"{submission.total_score if submission.total_score is not None else '-'} / {submission.max_score}"),
            ("得分率", f"{round((submission.score_rate or 0) * 100)}%" if submission.score_rate is not None else "未评估"),
            ("问卷完成时间", _format_datetime(submission.submitted_at)),
        ]
        result_table = document.add_table(rows=0, cols=2)
        result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        result_table.style = "Table Grid"
        for index in range(0, len(result), 2):
            row = result_table.add_row().cells
            for cell_index, item in enumerate(result[index:index + 2]):
                _set_cell_text(row[cell_index], *item)

        if submission.dimension_scores:
            dimension_title = document.add_paragraph()
            dimension_run = dimension_title.add_run("维度得分")
            dimension_run.bold = True
            _set_run_font(dimension_run)
            score_table = document.add_table(rows=1, cols=3)
            score_table.style = "Table Grid"
            score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ["诊断维度", "得分", "得分率"]
            for index, header in enumerate(headers):
                _set_cell_text(score_table.rows[0].cells[index], header, "")
            for item in sorted(submission.dimension_scores, key=lambda score: score.module.sort_order):
                row = score_table.add_row().cells
                values = [
                    _text(item.module.name),
                    f"{item.raw_score}/{item.max_score}",
                    f"{round(item.score_rate * 100)}%",
                ]
                for index, value in enumerate(values):
                    _set_plain_cell_text(row[index], value)

    _add_section_title(document, "三、AI 分析报告")
    for paragraph_text in _report_text(report.html_content if report else None).split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if paragraph_text:
            paragraph = document.add_paragraph(paragraph_text)
            _set_paragraph_font(paragraph)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
