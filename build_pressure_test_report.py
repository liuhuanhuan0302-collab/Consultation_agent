from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/Consultation_Agent_100并发压力测试报告.docx")
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "217346"
AMBER = "9A6700"
RED = "B42318"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_keep(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    ppr.append(keep)


def font_run(run, size=10.5, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, text="", size=10.5, bold=False, color="000000", after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    font_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, color="000000"):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    font_run(p.add_run(text), size=10.5, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def set_col_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def make_table(doc, headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run(str(text)), size=9.5, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            font_run(p.add_run(str(value)), size=9.2)
    set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.4)
sec.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for level, size, before, after in ((1, 16, 14, 7), (2, 13, 11, 5), (3, 11.5, 8, 4)):
    st = styles[f"Heading {level}"]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else NAVY)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font_run(hp.add_run("Consultation Agent | 性能测试报告"), size=8.5, color=MID_GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(fp.add_run("内部测试资料 | 2026-07-12"), size=8.5, color=MID_GRAY)

# Memo masthead
add_p(doc, "性能测试报告", size=10, bold=True, color=BLUE, after=4)
p = add_p(doc, "Consultation Agent\n100 并发压力测试报告", size=24, bold=True, color=NAVY, after=10)
p.paragraph_format.line_spacing = 1.0
add_p(doc, "基础业务链路 · 公网环境 · 阶梯并发验证", size=12.5, color=MID_GRAY, after=16)

meta = make_table(
    doc,
    ["项目", "内容"],
    [
        ("测试日期", "2026年7月12日"),
        ("测试对象", "Consultation Agent 公网部署环境"),
        ("公网地址", "http://8.138.165.2"),
        ("测试类型", "基础链路并发测试（不触发 AI 报告与邮件）"),
        ("测试工具", "项目内置 asyncio + httpx 压测脚本"),
    ],
    [1800, 7560],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "执行结论", 1)
cp = doc.add_paragraph()
cp.paragraph_format.left_indent = Inches(0.12)
cp.paragraph_format.right_indent = Inches(0.12)
cp.paragraph_format.space_before = Pt(5)
cp.paragraph_format.space_after = Pt(10)
cp.paragraph_format.line_spacing = 1.15
cp_pr = cp._p.get_or_add_pPr()
cp_shd = OxmlElement("w:shd")
cp_shd.set(qn("w:fill"), LIGHT_BLUE)
cp_pr.append(cp_shd)
font_run(cp.add_run("结论：本次 100 并发基础链路测试共发起 800 个请求，成功率 100%，压测后健康检查与数据库均正常；但各核心接口 P95 为 2.94–3.65 秒，未达到 P95 < 1.5 秒的建议目标，因此可判定“功能承载通过、性能目标未通过”。"), size=11, bold=True, color=NAVY)
doc.add_paragraph()

add_heading(doc, "关键指标", 2)
make_table(
    doc,
    ["指标", "结果", "判定"],
    [
        ("100 并发请求成功率", "100%（800/800）", "通过"),
        ("整体吞吐量", "37.25 req/s", "记录项"),
        ("接口 P95", "2.94–3.65 秒", "未达标"),
        ("压测后健康检查", "HTTP 200，database=ok", "通过"),
        ("完整报告链路", "未执行", "不在本次范围"),
    ],
    [3600, 3000, 2760],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
)

doc.add_page_break()
add_heading(doc, "1. 测试目标与范围", 1)
add_p(doc, "本次测试用于验证公网部署环境在阶梯并发负载下，基础用户链路的可用性、响应时间与稳定性，并重点确认系统是否能够完成一次 100 并发的短时突发测试。")
add_heading(doc, "1.1 覆盖链路", 2)
for item in [
    "创建匿名访问会话（POST /api/public/sessions）",
    "获取公开问卷题库（GET /api/public/questions）",
    "提交测试客户线索（POST /api/public/leads）",
    "保存部分问卷草稿（PUT /api/public/submissions/{id}/draft）",
]:
    add_bullet(doc, item)
add_heading(doc, "1.2 未覆盖范围", 2)
for item in [
    "完整 68 题提交、评分与报告生成",
    "DeepSeek 大模型调用及其限流、超时与费用影响",
    "PDF 浏览器渲染、邮件发送和 SMTP 服务能力",
    "持续 5–10 分钟的稳态负载、峰值后恢复和容量上限测试",
]:
    add_bullet(doc, item, color=MID_GRAY)

add_heading(doc, "2. 测试环境", 1)
make_table(
    doc,
    ["组件", "配置/状态"],
    [
        ("部署形态", "阿里云 ECS，Docker Compose"),
        ("入口", "公网 HTTP 80，经前端 Nginx 反向代理"),
        ("应用组件", "frontend、backend、report_worker、MySQL 8.0"),
        ("数据库", "MySQL 容器，健康状态正常"),
        ("连接池配置", "DB_POOL_SIZE=20；DB_MAX_OVERFLOW=40；DB_POOL_TIMEOUT=30"),
        ("健康接口", "/api/health"),
        ("环境标识", "健康接口显示 environment=development（建议改为 production）"),
    ],
    [2300, 7060],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "3. 测试方法", 1)
add_p(doc, "测试采用项目内置异步 HTTP 压测脚本，从独立客户端经公网访问服务器，按 10、30、100 并发逐级加压。每个虚拟用户顺序执行 4 个基础接口，测试数据使用 source_code=load_test 标记。")
make_table(
    doc,
    ["阶段", "模拟用户", "并发数", "理论请求数", "目的"],
    [
        ("阶段一", "20", "10", "80", "基础可用性与预热"),
        ("阶段二", "60", "30", "约 240", "中等并发稳定性"),
        ("阶段三", "200", "100", "800", "目标并发验证"),
    ],
    [1500, 1500, 1400, 1800, 3160],
    [WD_ALIGN_PARAGRAPH.CENTER] * 4 + [WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_page_break()
add_heading(doc, "4. 测试结果", 1)
add_heading(doc, "4.1 阶梯测试汇总", 2)
make_table(
    doc,
    ["并发", "完成请求", "成功情况", "吞吐量", "主要观察"],
    [
        ("10", "80", "100%", "7.39 req/s", "创建会话 P95 3.90 秒，响应偏慢"),
        ("30", "237", "236 成功、1 失败", "23.02 req/s", "出现 1 次 ReadError；创建会话成功率约 98.3%"),
        ("100", "800", "100%", "37.25 req/s", "零失败，但 P95 约 3 秒以上"),
    ],
    [1100, 1450, 1900, 1750, 3160],
    [WD_ALIGN_PARAGRAPH.CENTER] * 4 + [WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "4.2 100 并发接口明细", 2)
make_table(
    doc,
    ["接口", "次数", "成功率", "平均响应", "P95", "最大响应"],
    [
        ("创建会话", "200", "100%", "2267 ms", "3649 ms", "5907 ms"),
        ("获取题库", "200", "100%", "2136 ms", "2945 ms", "6458 ms"),
        ("提交线索", "200", "100%", "2402 ms", "3487 ms", "3901 ms"),
        ("保存草稿", "200", "100%", "2264 ms", "3021 ms", "4001 ms"),
    ],
    [2100, 1100, 1300, 1620, 1620, 1620],
    [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5,
)

add_heading(doc, "4.3 验收判断", 2)
make_table(
    doc,
    ["验收项", "目标", "实际", "结论"],
    [
        ("成功率", ">=99%", "100 并发阶段 100%", "通过"),
        ("P95 响应时间", "<1500 ms", "2945–3649 ms", "未通过"),
        ("服务错误", "无大量 500/502", "无 HTTP 失败", "通过"),
        ("数据库可用性", "压测后正常", "database=ok", "通过"),
        ("连续稳定性", "持续测试无退化", "本次未覆盖", "待验证"),
    ],
    [2600, 2100, 2780, 1880],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
)

add_heading(doc, "5. 风险与问题", 1)
for item in [
    "100 并发虽然零失败，但 P95 超出目标约 1.4–2.4 倍，活动高峰期可能出现明显等待感。",
    "30 并发阶段出现一次 ReadError，说明仍存在瞬时连接或网络抖动风险；单次 100 并发零失败不足以证明长期稳定。",
    "健康接口显示 environment=development，生产环境标识尚未规范化。",
    "完整报告链路涉及大模型、队列、Chromium PDF 与 SMTP，资源消耗和外部服务限流远高于本次基础链路。",
    "测试产生了标记为 load_test 的测试线索，需要在确认外键关系后清理。",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "6. 优化建议", 1)
make_table(
    doc,
    ["优先级", "建议", "预期目的"],
    [
        ("P0", "将 ENVIRONMENT 设置为 production，并保持现有数据库连接池配置", "规范生产配置，避免开发默认行为"),
        ("P0", "查看 ECS CPU、内存、磁盘 IO、网络带宽及 MySQL 慢查询", "定位 2–4 秒响应的主要瓶颈"),
        ("P1", "根据 CPU 核数增加 Uvicorn worker，并复测连接池与 MySQL 最大连接数", "提升并行请求处理能力"),
        ("P1", "对公开题库接口增加短期缓存，减少重复数据库查询", "降低高并发读压力"),
        ("P1", "执行 100 并发、持续 5–10 分钟的稳态压测", "验证错误率、资源增长和峰值后恢复"),
        ("P2", "单独进行 3→5→10 并发完整报告链路测试", "评估 DeepSeek、Worker、PDF 与邮件瓶颈"),
        ("P2", "配置域名、HTTPS 与生产监控告警", "提升安全性与可运营性"),
    ],
    [1100, 5360, 2900],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "7. 建议复测标准", 1)
for item in [
    "基础链路：100 并发持续 10 分钟，成功率 >=99%，无持续性 500/502/数据库连接超时。",
    "体验目标：核心接口 P95 <1500 ms；如业务接受更宽松目标，应由产品负责人书面确认。",
    "资源目标：CPU、内存、数据库连接数和磁盘 IO 不持续触顶，压测停止后可快速恢复。",
    "完整链路：先关闭或隔离真实模型费用，再按 3、5、10 并发验证报告生成成功率与队列积压。",
]:
    add_bullet(doc, item)

add_heading(doc, "附录 A：本次执行命令", 1)
commands = [
    r"python backend/scripts/load_test.py --host http://8.138.165.2 --users 20 --concurrency 10",
    r"python backend/scripts/load_test.py --host http://8.138.165.2 --users 60 --concurrency 30",
    r"python backend/scripts/load_test.py --host http://8.138.165.2 --users 200 --concurrency 100",
]
for cmd in commands:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    set_cell = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    set_cell.append(shd)
    run = p.add_run(cmd)
    font_run(run, size=8.8, name="Consolas", color=NAVY)

add_heading(doc, "附录 B：数据说明", 1)
add_p(doc, "本报告数据来自 2026年7月12日实际公网压测输出。测试结果仅代表当时服务器配置、网络条件和短时负载，不应直接等同于长期 SLA 或完整报告链路容量。")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
